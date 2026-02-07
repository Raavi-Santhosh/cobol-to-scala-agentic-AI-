"""DOCX generation from structured content (sections with title + body)."""
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def write_docx(
    sections: list[dict],
    output_path: str | Path,
    title: str | None = None,
) -> Path:
    """
    Create a DOCX from a list of sections.
    Each section: {"title": str, "body": str} or {"title": str, "paragraphs": list[str]}
    """
    doc = Document()
    if title:
        t = doc.add_heading(title, level=0)
        t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for sec in sections:
        heading = sec.get("title", "")
        body = sec.get("body")
        paragraphs = sec.get("paragraphs")
        if heading:
            doc.add_heading(heading, level=1)
        if body:
            for para in body.strip().split("\n\n"):
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.space_after = Pt(6)
        if paragraphs:
            for para in paragraphs:
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.space_after = Pt(6)

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path
